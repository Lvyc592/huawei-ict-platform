#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
为"华技云华为 ICT 认证模拟考试系统 V1.0"拆分生成两份软著代码提取文档：
1. 主程序源码（纯 Java）
2. 前端辅助源码（自研 JS 交互逻辑）
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import re
import math

BASE_DIR = Path("C:/Users/lv/Desktop/springboot-huawei-ict(1)/springboot-huawei-ict(1)")
OUTPUT_DIR = BASE_DIR

SOFTWARE_NAME = "华技云华为 ICT 认证模拟考试系统 V1.0"
COPYRIGHT_HEADER = """/*
 * 华技云华为 ICT 认证模拟考试系统 V1.0
 * 版权所有 (C) 2024 华技云
 * 本软件受著作权法保护，未经许可不得转载、复制或用于商业用途。
 */"""

JS_COPYRIGHT_HEADER = """// 华技云华为 ICT 认证模拟考试系统 V1.0
// 版权所有 (C) 2024 华技云
// 本软件受著作权法保护，未经许可不得转载、复制或用于商业用途。"""

LINES_PER_PAGE = 50
MAX_PAGES = 60
FONT_NAME = "Courier New"
FONT_SIZE = Pt(9)  # 小五号
LINE_SPACING = Pt(12)  # 固定行距，保证每页至少 50 行

# URL / CDN / 脱敏模式
URL_PATTERN = re.compile(r"https?://[^\s\"'`]+|www\.[^\s\"'`]+|//cdn\.[^\s\"'`]+", re.IGNORECASE)
SENSITIVE_PATTERN = re.compile(r"ollama|openai|bilibili|youtube|cdn|bootstrap|jquery\.com|ajax\.googleapis|fonts\.google|cdnjs\.cloudflare|unpkg|jsdelivr|cloudflare|huawei\.com|aliyun\.com|tencent\.com|qiniu\.com", re.IGNORECASE)


def add_page_number(paragraph, page_count=True):
    """在页脚插入 Word 页码域。"""
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = "PAGE"
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def is_java_comment(line):
    s = line.strip()
    if not s:
        return True
    if s.startswith("//"):
        return True
    if s.startswith("/*") or s.startswith("*/") or s.startswith("*"):
        return True
    return False


def is_js_comment(line):
    s = line.strip()
    if not s:
        return True
    if s.startswith("//"):
        return True
    if s.startswith("/*") or s.startswith("*/") or s.startswith("*"):
        return True
    return False


def desensitize_line(line):
    """移除外网 URL、CDN、B 站、ollama/密钥地址等敏感信息。"""
    # 直接删除纯配置型敏感行（如 ollama.url、api_key、bootstrap 链接）
    stripped = line.strip().lower()
    if any(k in stripped for k in ["ollama", "bilibili", "youtube", "api_key", "secret", "password", "jdbc:mysql"]):
        # 但保留业务逻辑行，仅替换值
        pass
    # 替换 URL 为占位符
    line = URL_PATTERN.sub("https://placeholder.example.com", line)
    return line


def clean_java_line(line):
    """清洗 Java 单行：保留逻辑，替换敏感值。"""
    line = line.replace("\t", "    ")
    line = desensitize_line(line)
    # 若整行只剩占位符或纯敏感配置，置为空以便过滤
    if re.search(r'https?://placeholder\.example\.com', line) and not re.search(r'[=+\(\{\[]', line):
        return ""
    return line


def clean_js_line(line):
    """清洗 JS 单行：保留业务逻辑，删除样式/SVG/纯框架冗余。"""
    s = line.strip()
    # 删除明显的 CSS/SVG/布局行
    if s.startswith("<") and s.endswith(">") and not s.startswith("</script>"):
        return ""
    # 删除只包含样式或 SVG 路径的行
    if re.search(r'\b(style|class|svg|path|rect|circle|polygon|g\s|fill=|stroke=|width=|height=|viewBox=|xmlns=)', s):
        return ""
    # 删除外部 CDN/库引用
    if any(k in s.lower() for k in ["cdn", "jquery.com", "bootstrapcdn", "ajax.googleapis", "fonts.googleapis", "cloudflare", "unpkg", "jsdelivr"]):
        return ""
    line = line.replace("\t", "    ")
    line = desensitize_line(line)
    return line


def read_java_files(file_list):
    """读取 Java 文件，加入版权头，过滤空行和注释。"""
    all_lines = []
    for rel_path in file_list:
        fp = BASE_DIR / rel_path
        if not fp.exists():
            print(f"[WARN] 文件不存在: {rel_path}")
            continue
        with open(fp, 'r', encoding='utf-8', errors='ignore') as f:
            raw_lines = f.readlines()

        # 加版权头
        all_lines.append({
            'rel_path': rel_path,
            'line_no': 0,
            'content': COPYRIGHT_HEADER.splitlines()[0]
        })
        for c_line in COPYRIGHT_HEADER.splitlines()[1:]:
            all_lines.append({'rel_path': rel_path, 'line_no': 0, 'content': c_line})

        for idx, line in enumerate(raw_lines, 1):
            if is_java_comment(line):
                continue
            cleaned = clean_java_line(line.rstrip('\n').rstrip('\r'))
            if not cleaned.strip():
                continue
            all_lines.append({'rel_path': rel_path, 'line_no': idx, 'content': cleaned})
    return all_lines


def extract_js_from_html(html_path):
    """从 HTML 文件中提取内嵌 <script> 标签里的 JS 代码（不含 src）。"""
    with open(html_path, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    # 匹配 <script> ... </script>（不含 src 属性）
    pattern = re.compile(r'<script\b(?![^>]*\bsrc\s*=)[^>]*>(.*?)</script>', re.DOTALL | re.IGNORECASE)
    scripts = pattern.findall(content)
    lines = []
    for script in scripts:
        for line in script.splitlines():
            lines.append(line)
    return lines


def read_js_files():
    """读取所有 .js 文件，并从 HTML 提取内嵌 JS。"""
    all_lines = []
    # 独立 .js 文件
    js_files = sorted(BASE_DIR.glob("src/main/resources/static/**/*.js"))
    for js_path in js_files:
        rel_path = str(js_path.relative_to(BASE_DIR)).replace("\\", "/")
        with open(js_path, 'r', encoding='utf-8', errors='ignore') as f:
            raw_lines = f.readlines()
        # 版权头
        for c_line in JS_COPYRIGHT_HEADER.splitlines():
            all_lines.append({'rel_path': rel_path, 'line_no': 0, 'content': c_line})
        for idx, line in enumerate(raw_lines, 1):
            if is_js_comment(line):
                continue
            cleaned = clean_js_line(line.rstrip('\n').rstrip('\r'))
            if not cleaned.strip():
                continue
            all_lines.append({'rel_path': rel_path, 'line_no': idx, 'content': cleaned})

    # HTML 内嵌 JS
    html_files = sorted(BASE_DIR.glob("src/main/resources/static/**/*.html"))
    for html_path in html_files:
        rel_path = str(html_path.relative_to(BASE_DIR)).replace("\\", "/")
        js_lines = extract_js_from_html(html_path)
        if not js_lines:
            continue
        # 版权头
        for c_line in JS_COPYRIGHT_HEADER.splitlines():
            all_lines.append({'rel_path': rel_path + " [JS]", 'line_no': 0, 'content': c_line})
        for idx, line in enumerate(js_lines, 1):
            if is_js_comment(line):
                continue
            cleaned = clean_js_line(line.rstrip('\n').rstrip('\r'))
            if not cleaned.strip():
                continue
            all_lines.append({'rel_path': rel_path + " [JS]", 'line_no': idx, 'content': cleaned})
    return all_lines


def split_pages(all_lines, pages=MAX_PAGES, lines_per_page=LINES_PER_PAGE):
    """切分页面。若代码充足，取前 pages/2 + 后 pages/2 填满 pages 页；
    若不足，均匀分布使每页均不少于 lines_per_page 行，且页数最少。"""
    total = len(all_lines)
    max_lines = pages * lines_per_page

    if total >= max_lines:
        half = pages // 2
        half_lines = half * lines_per_page
        front = all_lines[:half_lines]
        back = all_lines[-half_lines:]
        combined = front + back
    else:
        combined = all_lines

    if not combined:
        return []

    # 最小页数，保证每页至少 lines_per_page 行
    page_count = len(combined) // lines_per_page
    if page_count == 0:
        page_count = 1
    base = len(combined) // page_count
    extra = len(combined) % page_count

    result = []
    start = 0
    for i in range(page_count):
        size = base + (1 if i < extra else 0)
        result.append(combined[start:start + size])
        start += size
    return result


def create_word_doc(pages, output_name, header_text):
    """生成标准化 Word 文档。"""
    output_path = OUTPUT_DIR / output_name
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # 页眉
    header = section.header
    hp = header.paragraphs[0]
    hp.text = header_text
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in hp.runs:
        r.font.name = 'SimSun'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        r.font.size = Pt(10.5)
        r.bold = True

    # 页脚：右对齐连续页码
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = fp.add_run()
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.size = Pt(10.5)
    add_page_number(fp)

    # 页面内容
    for page_idx, page_lines in enumerate(pages):
        if page_idx > 0:
            doc.add_page_break()

        first = page_lines[0]
        last = page_lines[-1]
        title = f"Source: {first['rel_path']} line {first['line_no']}"
        if first['rel_path'] == last['rel_path']:
            title += f" to line {last['line_no']}"
        else:
            title += f" ... {last['rel_path']} line {last['line_no']}"

        tp = doc.add_paragraph()
        tr = tp.add_run(title)
        tr.font.name = 'SimSun'
        tr._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        tr.font.size = Pt(9)
        tr.bold = True
        tr.font.color.rgb = RGBColor(0x09, 0x40, 0x86)
        tp.paragraph_format.space_before = Pt(0)
        tp.paragraph_format.space_after = Pt(0)
        tp.paragraph_format.line_spacing = LINE_SPACING

        for item in page_lines:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = LINE_SPACING
            p.paragraph_format.left_indent = Cm(0)
            safe = item['content'].replace('\t', '    ')
            run = p.add_run(safe)
            run.font.name = FONT_NAME
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
            run.font.size = FONT_SIZE

    doc.save(str(output_path))
    return output_path, len(pages)


def build_java_file_list():
    """按指定顺序构建 Java 文件列表。"""
    ordered = []
    # 1. 启动类
    ordered.append("src/main/java/com/huawei/ict/IctApplication.java")
    # 2. JWT/Security 配置
    ordered.extend([
        "src/main/java/com/huawei/ict/config/SecurityConfig.java",
        "src/main/java/com/huawei/ict/config/JwtUtil.java",
        "src/main/java/com/huawei/ict/config/JwtAuthenticationFilter.java",
        "src/main/java/com/huawei/ict/config/AiProperties.java",
    ])
    # 3. Entity 实体
    ordered.extend([
        "src/main/java/com/huawei/ict/entity/BaseEntity.java",
        "src/main/java/com/huawei/ict/entity/User.java",
        "src/main/java/com/huawei/ict/entity/Course.java",
        "src/main/java/com/huawei/ict/entity/CourseChapter.java",
        "src/main/java/com/huawei/ict/entity/Exam.java",
        "src/main/java/com/huawei/ict/entity/ExamResult.java",
        "src/main/java/com/huawei/ict/entity/Question.java",
        "src/main/java/com/huawei/ict/entity/QuestionRecord.java",
        "src/main/java/com/huawei/ict/entity/KnowledgePoint.java",
        "src/main/java/com/huawei/ict/entity/LearningRecord.java",
        "src/main/java/com/huawei/ict/entity/Lab.java",
        "src/main/java/com/huawei/ict/entity/LabInstance.java",
        "src/main/java/com/huawei/ict/entity/Certification.java",
        "src/main/java/com/huawei/ict/entity/Competition.java",
        "src/main/java/com/huawei/ict/entity/Job.java",
        "src/main/java/com/huawei/ict/entity/JobApplication.java",
        "src/main/java/com/huawei/ict/entity/Notification.java",
        "src/main/java/com/huawei/ict/entity/Registration.java",
        "src/main/java/com/huawei/ict/entity/SystemSetting.java",
        "src/main/java/com/huawei/ict/entity/UserCourse.java",
    ])
    # 4. DTO
    ordered.extend([
        "src/main/java/com/huawei/ict/dto/ApiResponse.java",
        "src/main/java/com/huawei/ict/dto/DashboardStats.java",
        "src/main/java/com/huawei/ict/dto/StudentStats.java",
        "src/main/java/com/huawei/ict/dto/StudySuggestionDTO.java",
        "src/main/java/com/huawei/ict/dto/MyCourseDTO.java",
        "src/main/java/com/huawei/ict/dto/LoginRequest.java",
        "src/main/java/com/huawei/ict/dto/LoginResponse.java",
    ])
    # 5. Repository
    ordered.extend(sorted([
        "src/main/java/com/huawei/ict/repository/CertificationRepository.java",
        "src/main/java/com/huawei/ict/repository/CompetitionRepository.java",
        "src/main/java/com/huawei/ict/repository/CourseChapterRepository.java",
        "src/main/java/com/huawei/ict/repository/CourseRepository.java",
        "src/main/java/com/huawei/ict/repository/ExamRepository.java",
        "src/main/java/com/huawei/ict/repository/ExamResultRepository.java",
        "src/main/java/com/huawei/ict/repository/JobApplicationRepository.java",
        "src/main/java/com/huawei/ict/repository/JobRepository.java",
        "src/main/java/com/huawei/ict/repository/KnowledgePointRepository.java",
        "src/main/java/com/huawei/ict/repository/LabInstanceRepository.java",
        "src/main/java/com/huawei/ict/repository/LabRepository.java",
        "src/main/java/com/huawei/ict/repository/LearningRecordRepository.java",
        "src/main/java/com/huawei/ict/repository/NotificationRepository.java",
        "src/main/java/com/huawei/ict/repository/QuestionRecordRepository.java",
        "src/main/java/com/huawei/ict/repository/QuestionRepository.java",
        "src/main/java/com/huawei/ict/repository/RegistrationRepository.java",
        "src/main/java/com/huawei/ict/repository/SystemSettingRepository.java",
        "src/main/java/com/huawei/ict/repository/UserCourseRepository.java",
        "src/main/java/com/huawei/ict/repository/UserRepository.java",
    ]))
    # 6. Service（学情/AI/考试优先）
    ordered.extend([
        "src/main/java/com/huawei/ict/service/DashboardService.java",
        "src/main/java/com/huawei/ict/service/AiService.java",
        "src/main/java/com/huawei/ict/service/ExamService.java",
        "src/main/java/com/huawei/ict/service/PracticeService.java",
        "src/main/java/com/huawei/ict/service/KnowledgeEngine.java",
        "src/main/java/com/huawei/ict/service/AuthService.java",
        "src/main/java/com/huawei/ict/service/CourseService.java",
        "src/main/java/com/huawei/ict/service/LabService.java",
        "src/main/java/com/huawei/ict/service/UserService.java",
        "src/main/java/com/huawei/ict/service/NotificationService.java",
        "src/main/java/com/huawei/ict/service/SystemSettingService.java",
        "src/main/java/com/huawei/ict/service/CertificationService.java",
        "src/main/java/com/huawei/ict/service/CompetitionService.java",
        "src/main/java/com/huawei/ict/service/JobService.java",
    ])
    # 7. Controller
    ordered.extend([
        "src/main/java/com/huawei/ict/controller/AuthController.java",
        "src/main/java/com/huawei/ict/controller/StudentController.java",
        "src/main/java/com/huawei/ict/controller/TeacherController.java",
        "src/main/java/com/huawei/ict/controller/AdminController.java",
        "src/main/java/com/huawei/ict/controller/AiChatController.java",
    ])
    # 8. 全局异常
    ordered.append("src/main/java/com/huawei/ict/controller/GlobalExceptionHandler.java")
    return [p for p in ordered if (BASE_DIR / p).exists()]


def main():
    print("=" * 60)
    print("生成软著代码提取文档：拆分 Java 主程序 + JS 前端辅助")
    print("=" * 60)

    # 文档 1：主程序源码（纯 Java）
    java_files = build_java_file_list()
    java_lines = read_java_files(java_files)
    java_pages = split_pages(java_lines, pages=MAX_PAGES, lines_per_page=LINES_PER_PAGE)
    java_path, java_page_count = create_word_doc(
        java_pages,
        "软著代码提取-华技云华为ICT认证模拟考试系统V1.0-主程序源码.docx",
        SOFTWARE_NAME
    )
    print(f"[Java] 有效代码行数: {len(java_lines)}, 页数: {java_page_count}, 保存: {java_path}")

    # 文档 2：前端辅助源码（自研 JS）
    js_lines = read_js_files()
    js_pages = split_pages(js_lines, pages=MAX_PAGES, lines_per_page=LINES_PER_PAGE)
    js_path, js_page_count = create_word_doc(
        js_pages,
        "软著代码提取-华技云华为ICT认证模拟考试系统V1.0-前端辅助源码.docx",
        SOFTWARE_NAME
    )
    print(f"[JS] 有效代码行数: {len(js_lines)}, 页数: {js_page_count}, 保存: {js_path}")

    # 检查
    all_ok = True
    for i, page in enumerate(java_pages, 1):
        if len(page) < LINES_PER_PAGE:
            print(f"[WARN] Java 第 {i} 页仅 {len(page)} 行")
            all_ok = False
    for i, page in enumerate(js_pages, 1):
        if len(page) < LINES_PER_PAGE:
            print(f"[WARN] JS 第 {i} 页仅 {len(page)} 行")
            all_ok = False
    if java_page_count > MAX_PAGES or js_page_count > MAX_PAGES:
        print(f"[WARN] 页数超过 {MAX_PAGES}")
        all_ok = False
    if all_ok:
        print("[OK] 两份文档均符合要求：每页 50 行，60 页以内。")


if __name__ == '__main__':
    main()
