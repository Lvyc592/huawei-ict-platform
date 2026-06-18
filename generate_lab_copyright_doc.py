#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
generate copyright doc for lab tracking system
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

BASE_DIR = Path("C:/Users/lv/Desktop/springboot-huawei-ict(1)/springboot-huawei-ict(1)")
OUTPUT_PATH = BASE_DIR / "软著代码提取-华技云华为云实训过程全追踪管理系统V1.0.docx"

SOFTWARE_NAME = "华技云华为云实训过程全追踪管理系统 V1.0"
LINES_PER_PAGE = 50
MAX_TOTAL_PAGES = 60
LINE_FONT_SIZE = Pt(8)
LINE_SPACING = Pt(11.5)

BACKEND_FILES = [
    "src/main/java/com/huawei/ict/IctApplication.java",
    "src/main/java/com/huawei/ict/config/SecurityConfig.java",
    "src/main/java/com/huawei/ict/config/JwtUtil.java",
    "src/main/java/com/huawei/ict/config/JwtAuthenticationFilter.java",
    "src/main/java/com/huawei/ict/config/AiProperties.java",
    "src/main/java/com/huawei/ict/dto/ApiResponse.java",
    "src/main/java/com/huawei/ict/dto/MyCourseDTO.java",
    "src/main/java/com/huawei/ict/dto/StudySuggestionDTO.java",
    "src/main/java/com/huawei/ict/dto/DashboardStats.java",
    "src/main/java/com/huawei/ict/entity/Lab.java",
    "src/main/java/com/huawei/ict/entity/LabInstance.java",
    "src/main/java/com/huawei/ict/entity/User.java",
    "src/main/java/com/huawei/ict/entity/Course.java",
    "src/main/java/com/huawei/ict/entity/CourseChapter.java",
    "src/main/java/com/huawei/ict/entity/Certification.java",
    "src/main/java/com/huawei/ict/entity/LearningRecord.java",
    "src/main/java/com/huawei/ict/entity/Notification.java",
    "src/main/java/com/huawei/ict/repository/LabRepository.java",
    "src/main/java/com/huawei/ict/repository/LabInstanceRepository.java",
    "src/main/java/com/huawei/ict/repository/UserRepository.java",
    "src/main/java/com/huawei/ict/repository/CourseRepository.java",
    "src/main/java/com/huawei/ict/service/LabService.java",
    "src/main/java/com/huawei/ict/service/DashboardService.java",
    "src/main/java/com/huawei/ict/service/CourseService.java",
    "src/main/java/com/huawei/ict/service/AuthService.java",
    "src/main/java/com/huawei/ict/service/NotificationService.java",
    "src/main/java/com/huawei/ict/service/AiService.java",
    "src/main/java/com/huawei/ict/controller/AdminController.java",
    "src/main/java/com/huawei/ict/controller/StudentController.java",
    "src/main/java/com/huawei/ict/controller/TeacherController.java",
    "src/main/java/com/huawei/ict/controller/AuthController.java",
    "src/main/java/com/huawei/ict/controller/AiChatController.java",
    "src/main/java/com/huawei/ict/controller/GlobalExceptionHandler.java",
]

FRONTEND_FILES = [
    "src/main/resources/static/index.html",
    "src/main/resources/static/admin/dashboard.html",
    "src/main/resources/static/admin/labs.html",
    "src/main/resources/static/admin/courses.html",
    "src/main/resources/static/admin/users.html",
    "src/main/resources/static/admin/certifications.html",
    "src/main/resources/static/admin/analytics.html",
    "src/main/resources/static/admin/settings.html",
    "src/main/resources/static/teacher/dashboard.html",
    "src/main/resources/static/teacher/courses.html",
    "src/main/resources/static/teacher/students.html",
    "src/main/resources/static/teacher/analytics.html",
    "src/main/resources/static/teacher/exam-records.html",
    "src/main/resources/static/student/dashboard.html",
    "src/main/resources/static/student/cloud-lab.html",
    "src/main/resources/static/student/lab-env.html",
    "src/main/resources/static/student/lab-loading.html",
    "src/main/resources/static/student/courses.html",
    "src/main/resources/static/student/course-detail.html",
    "src/main/resources/static/student/ai-tutor.html",
    "src/main/resources/static/student/analytics.html",
]


def add_page_number(paragraph):
    run = paragraph.add_run()
    fc1 = OxmlElement("w:fldChar")
    fc1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText")
    it.set(qn("xml:space"), "preserve")
    it.text = "PAGE"
    fc2 = OxmlElement("w:fldChar")
    fc2.set(qn("w:fldCharType"), "end")
    run._r.append(fc1)
    run._r.append(it)
    run._r.append(fc2)


def is_bad_line(line):
    s = line.strip()
    if not s: return True
    if s.startswith("//") or s.startswith("#"): return True
    if s.startswith("/*") or s.startswith("*/"): return True
    if s.startswith("/**") or s.startswith("*"): return True
    if s.startswith("<!--") or s.startswith("-->"): return True
    return False


def read_valid(rel):
    fp = BASE_DIR / rel
    try:
        with open(fp, "r", encoding="utf-8", errors="ignore") as f:
            lines = f.readlines()
        out = []
        for line in lines:
            if not is_bad_line(line):
                out.append(line.rstrip("\n").rstrip("\r"))
        return out
    except Exception as e:
        print("read fail:", rel, e)
        return []


def collect(files):
    out = []
    for rel in files:
        ls = read_valid(rel)
        for idx, line in enumerate(ls, 1):
            out.append({"rel": rel, "no": idx, "txt": line})
    return out


def split_pages(all_lines, total=30):
    half = total // 2
    per_half = half * LINES_PER_PAGE
    if len(all_lines) <= per_half * 2:
        pages = []
        for i in range(0, len(all_lines), LINES_PER_PAGE):
            pages.append(all_lines[i:i+LINES_PER_PAGE])
        return pages
    front = all_lines[:per_half]
    back = all_lines[-per_half:]
    combined = front + back
    pages = []
    for i in range(0, len(combined), LINES_PER_PAGE):
        pages.append(combined[i:i+LINES_PER_PAGE])
    return pages


def make_doc(bp, fp):
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    # header
    hdr = sec.header
    hp = hdr.paragraphs[0]
    hp.text = SOFTWARE_NAME
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in hp.runs:
        r.font.name = "SimSun"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        r.font.size = Pt(10.5)
        r.bold = True
    # footer
    ftr = sec.footer
    fp_para = ftr.paragraphs[0]
    fp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp_para.add_run("Page ")
    add_page_number(fp_para)
    fp_para.add_run(" of")
    for r in fp_para.runs:
        r.font.name = "SimSun"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        r.font.size = Pt(10.5)
    all_pages = bp + fp
    for pi, plines in enumerate(all_pages):
        if pi > 0:
            doc.add_page_break()
        first = plines[0]
        last = plines[-1]
        title = "Source: {} line {}".format(first["rel"], first["no"])
        if first["rel"] == last["rel"]:
            title += " to {}".format(last["no"])
        else:
            title += " ... {} line {}".format(last["rel"], last["no"])
        tp = doc.add_paragraph()
        tr = tp.add_run(title)
        tr.font.name = "SimSun"
        tr._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        tr.font.size = Pt(9)
        tr.bold = True
        tr.font.color.rgb = RGBColor(0x09, 0x40, 0x86)
        tp.paragraph_format.space_before = Pt(0)
        tp.paragraph_format.space_after = Pt(0)
        tp.paragraph_format.line_spacing = LINE_SPACING
        for item in plines:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = LINE_SPACING
            p.paragraph_format.left_indent = Cm(0)
            safe = item["txt"].replace("\t", "    ")
            run = p.add_run(safe)
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
            run.font.size = LINE_FONT_SIZE
    doc.save(str(OUTPUT_PATH))
    return len(all_pages)


def main():
    print("=" * 60)
    print("generating copyright doc...")
    print("software:", SOFTWARE_NAME)
    print("=" * 60)
    bl = collect(BACKEND_FILES)
    fl = collect(FRONTEND_FILES)
    print("backend lines:", len(bl))
    print("frontend lines:", len(fl))
    print("total lines:", len(bl) + len(fl))
    bp = split_pages(bl, 30)
    fp = split_pages(fl, 30)
    print("backend pages:", len(bp))
    print("frontend pages:", len(fp))
    total = make_doc(bp, fp)
    print("saved:", OUTPUT_PATH)
    print("total pages:", total)
    ok = True
    for i, pg in enumerate(bp + fp, 1):
        if len(pg) < LINES_PER_PAGE:
            print("WARN: page", i, "only", len(pg), "lines")
            ok = False
    if total > MAX_TOTAL_PAGES:
        print("WARN: too many pages")
        ok = False
    if ok:
        print("OK: all requirements met")


if __name__ == "__main__":
    main()
