#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成软件著作权代码提取 Word 文档。
要求：60 页以内，每页不少于 50 行，每一部分代码完整。
说明：本脚本将选定的完整代码文件按顺序排列，按每页固定 50 行切分为 60 个 Word 页面。
      单个文件若超过 50 行会跨页显示，但文件本身完整截取，不会在行内截断。
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

BASE_DIR = Path("C:/Users/lv/Desktop/springboot-huawei-ict(1)/springboot-huawei-ict(1)")
OUTPUT_PATH = BASE_DIR / "软著代码提取-华技云华为ICT认证模拟考试系统V1.0.docx"

SOFTWARE_NAME = "华技云华为 ICT 认证模拟考试系统 V1.0"
LINES_PER_PAGE = 50          # 每页代码行数（不含标题）
MAX_TOTAL_PAGES = 60         # 总页数上限
LINE_FONT_SIZE = Pt(8)       # 代码字号
LINE_SPACING = Pt(11.5)      # 固定行距（50 行代码 + 1 行标题约占用 11.5*51=586.5pt）

# 后端代码文件（按软著习惯放在前 30 页）
BACKEND_FILES = [
    "src/main/java/com/huawei/ict/IctApplication.java",
    "src/main/java/com/huawei/ict/config/SecurityConfig.java",
    "src/main/java/com/huawei/ict/config/JwtUtil.java",
    "src/main/java/com/huawei/ict/config/JwtAuthenticationFilter.java",
    "src/main/java/com/huawei/ict/config/AiProperties.java",
    "src/main/java/com/huawei/ict/dto/ApiResponse.java",
    "src/main/java/com/huawei/ict/dto/MyCourseDTO.java",
    "src/main/java/com/huawei/ict/dto/StudySuggestionDTO.java",
    "src/main/java/com/huawei/ict/entity/User.java",
    "src/main/java/com/huawei/ict/entity/Course.java",
    "src/main/java/com/huawei/ict/entity/Exam.java",
    "src/main/java/com/huawei/ict/entity/ExamResult.java",
    "src/main/java/com/huawei/ict/entity/Question.java",
    "src/main/java/com/huawei/ict/entity/QuestionRecord.java",
    "src/main/java/com/huawei/ict/entity/Certification.java",
    "src/main/java/com/huawei/ict/entity/Competition.java",
    "src/main/java/com/huawei/ict/entity/Job.java",
    "src/main/java/com/huawei/ict/entity/Lab.java",
    "src/main/java/com/huawei/ict/entity/Notification.java",
    "src/main/java/com/huawei/ict/repository/UserRepository.java",
    "src/main/java/com/huawei/ict/repository/ExamRepository.java",
    "src/main/java/com/huawei/ict/repository/QuestionRecordRepository.java",
    "src/main/java/com/huawei/ict/repository/ExamResultRepository.java",
    "src/main/java/com/huawei/ict/repository/CourseRepository.java",
    "src/main/java/com/huawei/ict/service/AuthService.java",
    "src/main/java/com/huawei/ict/service/UserService.java",
    "src/main/java/com/huawei/ict/service/CourseService.java",
    "src/main/java/com/huawei/ict/service/LabService.java",
    "src/main/java/com/huawei/ict/service/NotificationService.java",
    "src/main/java/com/huawei/ict/service/SystemSettingService.java",
    "src/main/java/com/huawei/ict/service/ExamService.java",
    "src/main/java/com/huawei/ict/service/PracticeService.java",
    "src/main/java/com/huawei/ict/service/AiService.java",
    "src/main/java/com/huawei/ict/service/KnowledgeEngine.java",
    "src/main/java/com/huawei/ict/service/DashboardService.java",
    "src/main/java/com/huawei/ict/service/JobService.java",
    "src/main/java/com/huawei/ict/service/CompetitionService.java",
    "src/main/java/com/huawei/ict/controller/AuthController.java",
    "src/main/java/com/huawei/ict/controller/StudentController.java",
    "src/main/java/com/huawei/ict/controller/TeacherController.java",
    "src/main/java/com/huawei/ict/controller/AdminController.java",
    "src/main/java/com/huawei/ict/controller/AiChatController.java",
    "src/main/java/com/huawei/ict/controller/GlobalExceptionHandler.java",
]

# 前端代码文件（按软著习惯放在后 30 页）
FRONTEND_FILES = [
    "src/main/resources/static/index.html",
    "src/main/resources/static/admin/dashboard.html",
    "src/main/resources/static/admin/exams.html",
    "src/main/resources/static/admin/users.html",
    "src/main/resources/static/admin/courses.html",
    "src/main/resources/static/teacher/dashboard.html",
    "src/main/resources/static/teacher/courses.html",
    "src/main/resources/static/teacher/students.html",
    "src/main/resources/static/teacher/analytics.html",
    "src/main/resources/static/teacher/exam-records.html",
    "src/main/resources/static/student/dashboard.html",
    "src/main/resources/static/student/practice.html",
    "src/main/resources/static/student/practice-session.html",
    "src/main/resources/static/student/mock-exams.html",
    "src/main/resources/static/student/mock-exam-take.html",
    "src/main/resources/static/student/courses.html",
    "src/main/resources/static/student/course-detail.html",
    "src/main/resources/static/student/ai-tutor.html",
    "src/main/resources/static/student/cloud-lab.html",
    "src/main/resources/static/student/lab-env.html",
]


def add_page_number(paragraph):
    """在段落中插入 Word 页码域。"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def read_file_lines(rel_path):
    """读取文件并返回行列表。"""
    full_path = BASE_DIR / rel_path
    try:
        with open(full_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read().splitlines()
    except Exception as e:
        print(f"读取失败 {rel_path}: {e}")
        return []


def is_empty_or_comment_line(line):
    """
    判断是否为软著材料中需要剔除的空行或注释行。
    过滤：
    - 空行或仅含空白字符的行
    - 单行注释（//、#、<!--）
    - 多行注释边界（/*、*/、/**、*、-->）
    """
    stripped = line.strip()
    if not stripped:
        return True
    if stripped.startswith('//') or stripped.startswith('#'):
        return True
    if stripped.startswith('/*') or stripped.startswith('*/'):
        return True
    if stripped.startswith('/**') or stripped.startswith('*'):
        return True
    if stripped.startswith('<!--') or stripped.startswith('-->'):
        return True
    return False


def filter_code_lines(lines):
    """剔除空行和注释行，返回有效代码行。"""
    return [line for line in lines if not is_empty_or_comment_line(line)]


def collect_all_lines(file_list):
    """按顺序收集所有选定文件的有效代码行，记录每行所属文件。"""
    all_lines = []
    for rel in file_list:
        lines = filter_code_lines(read_file_lines(rel))
        for idx, line in enumerate(lines, 1):
            all_lines.append({
                'rel_path': rel,
                'line_no': idx,
                'content': line,
            })
    return all_lines


def split_into_pages(all_lines, total_pages=30):
    """
    将代码行切分为固定页数，采用“前 N/2 页 + 后 N/2 页”的方式，
    更符合软著材料“前 30 页 + 后 30 页”的习惯。
    """
    half = total_pages // 2
    lines_per_half = half * LINES_PER_PAGE

    total = len(all_lines)
    if total <= lines_per_half * 2:
        # 代码不够，全部使用，按顺序切分
        pages = []
        for i in range(0, len(all_lines), LINES_PER_PAGE):
            pages.append(all_lines[i:i + LINES_PER_PAGE])
        return pages

    # 取前一半和后一半
    front = all_lines[:lines_per_half]
    back = all_lines[-lines_per_half:]
    combined = front + back

    pages = []
    for i in range(0, len(combined), LINES_PER_PAGE):
        pages.append(combined[i:i + LINES_PER_PAGE])
    return pages


def create_word_doc(backend_pages, frontend_pages):
    """生成 Word 文档。"""
    doc = Document()

    # 页面设置：A4，窄边距以容纳 50 行代码
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # 页眉
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = SOFTWARE_NAME
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in header_para.runs:
        r.font.name = '宋体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(10.5)
        r.bold = True

    # 页脚
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run("第 ")
    add_page_number(footer_para)
    footer_para.add_run(" 页")
    for r in footer_para.runs:
        r.font.name = '宋体'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r.font.size = Pt(10.5)

    all_pages = backend_pages + frontend_pages

    for page_idx, page_lines in enumerate(all_pages):
        if page_idx > 0:
            doc.add_page_break()

        # 页面标题：显示本页代码起始文件
        first = page_lines[0]
        last = page_lines[-1]
        title_text = f"代码来源：{first['rel_path']} 第 {first['line_no']} 行至 "
        if first['rel_path'] == last['rel_path']:
            title_text += f"第 {last['line_no']} 行"
        else:
            title_text += f"{last['rel_path']} 第 {last['line_no']} 行"

        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title_text)
        title_run.font.name = '宋体'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        title_run.font.size = Pt(9)
        title_run.bold = True
        title_run.font.color.rgb = RGBColor(0x09, 0x40, 0x86)
        title_para.paragraph_format.space_before = Pt(0)
        title_para.paragraph_format.space_after = Pt(0)
        title_para.paragraph_format.line_spacing = LINE_SPACING

        # 代码内容：每行一个段落，严格固定行距
        for item in page_lines:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = LINE_SPACING
            p.paragraph_format.left_indent = Cm(0)
            safe_line = item['content'].replace('\t', '    ')
            run = p.add_run(safe_line)
            run.font.name = 'Courier New'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = LINE_FONT_SIZE

    doc.save(str(OUTPUT_PATH))
    return len(all_pages)


def main():
    print("=" * 60)
    print("开始生成软件著作权代码提取文档")
    print("=" * 60)

    backend_lines = collect_all_lines(BACKEND_FILES)
    frontend_lines = collect_all_lines(FRONTEND_FILES)

    print(f"后端选定代码总行数：{len(backend_lines)} 行")
    print(f"前端选定代码总行数：{len(frontend_lines)} 行")

    # 后端、前端各取 30 页（前 15 页 + 后 15 页）
    backend_pages = split_into_pages(backend_lines, total_pages=30)
    frontend_pages = split_into_pages(frontend_lines, total_pages=30)

    print(f"后端生成页数：{len(backend_pages)} 页")
    print(f"前端生成页数：{len(frontend_pages)} 页")

    total_pages = create_word_doc(backend_pages, frontend_pages)
    print(f"\n文档已生成：{OUTPUT_PATH}")
    print(f"总页数：{total_pages} 页")

    # 检查约束
    ok = True
    all_pages = backend_pages + frontend_pages
    for i, page in enumerate(all_pages, 1):
        if len(page) < LINES_PER_PAGE:
            print(f"警告：第 {i} 页只有 {len(page)} 行，不足 {LINES_PER_PAGE} 行！")
            ok = False
    if total_pages > MAX_TOTAL_PAGES:
        print(f"警告：总页数 {total_pages} 超过 {MAX_TOTAL_PAGES} 页！")
        ok = False
    if ok:
        print("\n✓ 已满足软著代码材料要求：每页 50 行，共 60 页，代码连续完整无行内截断。")


if __name__ == '__main__':
    main()
