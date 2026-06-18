#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成《华技云华为 ICT 认证模拟考试系统》用户/教师/管理员操作手册 Word 文档。"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

BASE_DIR = Path("C:/Users/lv/Desktop/springboot-huawei-ict(1)/springboot-huawei-ict(1)")
OUTPUT_PATH = BASE_DIR / "华技云华为ICT认证模拟考试系统-操作手册.docx"
DOC_TITLE = "华技云华为 ICT 认证模拟考试系统操作手册"
SOFTWARE_NAME = "华技云华为 ICT 认证模拟考试系统 V1.0"

def set_run_font(run, font_name='宋体', font_size=Pt(10.5), bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size
    run.bold = bold
    if color:
        run.font.color.rgb = color

def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    font_name = '黑体' if level == 1 else ('宋体' if level >= 3 else '楷体')
    font_size = {1: Pt(18), 2: Pt(15), 3: Pt(14), 4: Pt(12)}.get(level, Pt(12))
    set_run_font(run, font_name=font_name, font_size=font_size, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if level > 1 else WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_paragraph(doc, text, indent=True, bold=False, color=None):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, font_name='宋体', font_size=Pt(10.5), bold=bold, color=color)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, font_size=Pt(10.5))
    return p

def add_numbered(doc, text, level=0):
    p = doc.add_paragraph(style='List Number' if level == 0 else 'List Number 2')
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, font_size=Pt(10.5))
    return p

def add_table_intro(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                set_run_font(r, font_size=Pt(10.5), bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row):
            row_cells[i].text = str(cell_text)
            for p in row_cells[i].paragraphs:
                for r in p.runs:
                    set_run_font(r, font_size=Pt(10.5))
    doc.add_paragraph()
    return table

def add_page_number(section):
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("第 ")
    set_run_font(run, font_size=Pt(10.5))
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    set_run_font(run, font_size=Pt(10.5))
    run = footer_para.add_run(" 页")
    set_run_font(run, font_size=Pt(10.5))

def create_manual():
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 页眉
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = SOFTWARE_NAME
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in header_para.runs:
        set_run_font(r, font_size=Pt(10.5), bold=True)

    # 页脚
    add_page_number(section)

    # 封面
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(DOC_TITLE)
    set_run_font(run, font_name='黑体', font_size=Pt(26), bold=True, color=RGBColor(0x09, 0x40, 0x86))

    doc.add_paragraph()
    doc.add_paragraph()
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("用户 / 教师 / 管理员 使用指南")
    set_run_font(run, font_name='楷体', font_size=Pt(16))

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run("版本：V1.0\n编制日期：2026 年 6 月")
    set_run_font(run, font_name='宋体', font_size=Pt(12))

    doc.add_page_break()

    # 目录占位
    add_heading(doc, "目录", level=1)
    toc = [
        "1  文档说明",
        "2  系统概述",
        "3  通用操作（登录与退出）",
        "4  学生端操作手册",
        "   4.1  学生仪表盘",
        "   4.2  题库练习",
        "   4.3  模拟考试",
        "   4.4  课程学习",
        "   4.5  AI 辅导",
        "   4.6  云端实验",
        "   4.7  个人中心与认证",
        "5  教师端操作手册",
        "   5.1  教学仪表盘",
        "   5.2  我的课程",
        "   5.3  学员管理",
        "   5.4  考试记录",
        "   5.5  学情分析",
        "6  管理员端操作手册",
        "   6.1  管理仪表盘",
        "   6.2  用户管理",
        "   6.3  课程管理",
        "   6.4  认证管理",
        "   6.5  考试题库",
        "   6.6  实训管理",
        "   6.7  赛事管理",
        "   6.8  就业管理",
        "   6.9  数据统计",
        "   6.10 系统设置",
        "7  常见问题与注意事项",
    ]
    for item in toc:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_run_font(run, font_size=Pt(10.5))
    doc.add_page_break()

    # 1 文档说明
    add_heading(doc, "1  文档说明", level=1)
    add_paragraph(doc, "本手册面向《华技云华为 ICT 认证模拟考试系统》（以下简称“本系统”）的三类主要用户：学生、教师、系统管理员。手册按照角色使用场景分别介绍各功能模块的操作方法、界面说明及注意事项，帮助用户快速上手并规范使用系统。")
    add_paragraph(doc, "阅读对象：")
    add_bullet(doc, "学生：进行题库练习、模拟考试、课程学习、AI 辅导及云端实验等操作。")
    add_bullet(doc, "教师：管理所授课程、查看学员信息、查阅考试记录、分析学情数据。")
    add_bullet(doc, "管理员：负责系统用户、课程、题库、认证、实训、赛事、就业、统计及系统配置等运维管理工作。")
    doc.add_page_break()

    # 2 系统概述
    add_heading(doc, "2  系统概述", level=1)
    add_paragraph(doc, "本系统是一款面向华为 ICT 认证学习的在线模拟考试与教学管理平台，基于 B/S 架构开发，支持学生端、教师端、管理员端三类角色。系统主要功能包括：")
    add_bullet(doc, "题库练习：按认证方向与知识点分类，支持学生反复练习，答错后提供 AI 解析。")
    add_bullet(doc, "模拟考试：模拟真实华为 ICT 认证考试环境，支持限时作答、自动评分与成绩统计。")
    add_bullet(doc, "课程学习：提供课程资料、视频、实验指导等学习资源。")
    add_bullet(doc, "AI 辅导：基于大模型的智能答疑与错题解析。")
    add_bullet(doc, "云端实验：提供华为 ICT 相关实验环境，支持在线操作。")
    add_bullet(doc, "教师管理：教师可管理所授课程、学员、查看考试记录与学情分析。")
    add_bullet(doc, "后台管理：管理员进行用户、课程、题库、认证、实训、赛事、就业、统计等全平台管理。")
    doc.add_page_break()

    # 3 通用操作
    add_heading(doc, "3  通用操作（登录与退出）", level=1)
    add_heading(doc, "3.1  登录系统", level=2)
    add_numbered(doc, "打开浏览器，访问系统首页（index.html）。")
    add_numbered(doc, "在登录入口选择对应的角色入口：学生端、教师端或管理员端。")
    add_numbered(doc, "输入已注册的用户名（或学号/工号）和密码。")
    add_numbered(doc, "点击“登录”按钮，系统校验身份后进入对应角色的工作台。")
    add_numbered(doc, "若首次登录或忘记密码，可点击“忘记密码”按系统提示进行找回或联系管理员重置。")

    add_heading(doc, "3.2  退出系统", level=2)
    add_numbered(doc, "将鼠标移至页面右上角用户头像或“退出”按钮。")
    add_numbered(doc, "点击“退出登录”或“退出”，系统清除当前登录状态并返回登录首页。")
    add_numbered(doc, "为安全起见，建议离开电脑前主动退出。")
    doc.add_page_break()

    # 4 学生端
    add_heading(doc, "4  学生端操作手册", level=1)
    add_paragraph(doc, "学生端是系统的主要使用入口，提供刷题、考试、学习、AI 辅导、实验等功能。")

    add_heading(doc, "4.1  学生仪表盘", level=2)
    add_paragraph(doc, "登录后首先进入学生仪表盘。仪表盘汇总展示学习进度、练习数据、考试安排、课程推荐、最近学习记录等核心信息，帮助学生快速了解当前学习状态。")
    add_bullet(doc, "顶部统计卡片：显示学习天数、累计刷题、正确率、认证进度等关键指标。")
    add_bullet(doc, "快捷入口：可一键进入“题库练习”“模拟考试”“我的课程”“AI 辅导”。")
    add_bullet(doc, "今日任务：展示系统推荐的今日练习与课程学习任务。")
    add_bullet(doc, "消息通知：查看系统公告、课程更新、考试提醒等。")

    add_heading(doc, "4.2  题库练习", level=2)
    add_paragraph(doc, "题库练习模块用于按知识点或认证方向进行日常刷题，答错后可查看 AI 解析。")
    add_numbered(doc, "在左侧导航栏点击“题库练习”。")
    add_numbered(doc, "选择练习模式：可按科目、认证方向（如 HCIA、HCIP、HCIE）、知识点标签或难度筛选。")
    add_numbered(doc, "设置练习数量后点击“开始练习”，进入练习页面。")
    add_numbered(doc, "逐题作答，选择答案后点击“提交答案”。")
    add_numbered(doc, "系统即时反馈对错：答对显示正确提示；答错显示正确答案并触发“华小智 · AI 解析”，提供知识点定位、错因分析、正确解析与记忆技巧。")
    add_numbered(doc, "练习完成后可在练习报告页查看本次正确率、用时、错题分布及继续练习建议。")

    add_heading(doc, "4.3  模拟考试", level=2)
    add_paragraph(doc, "模拟考试模块用于全真模拟华为 ICT 认证考试，支持限时、计分、成绩统计。")
    add_numbered(doc, "在左侧导航栏点击“模拟考试”。")
    add_numbered(doc, "浏览可参加的考试列表，查看考试名称、认证方向、时长、及格分数、题目数量。")
    add_numbered(doc, "点击“开始考试”进入独立考试页面，倒计时开始。")
    add_numbered(doc, "按顺序作答，支持标记不确定的题目，便于后续检查。")
    add_numbered(doc, "作答完成后点击“提交试卷”。系统将在后台自动评分。")
    add_numbered(doc, "考试结束后可查看成绩报告，包括得分、合格状态、正确率、各知识点得分分布。")
    add_numbered(doc, "已完成的考试记录可在教师端“考试记录”中同步查看。")
    doc.add_page_break()

    add_heading(doc, "4.4  课程学习", level=2)
    add_paragraph(doc, "课程学习模块提供华为 ICT 相关课程资料与视频学习入口。")
    add_numbered(doc, "点击“课程中心”或“我的课程”。")
    add_numbered(doc, "浏览课程列表，可按认证方向、课程状态、任课教师筛选。")
    add_numbered(doc, "点击课程封面进入课程详情页，查看课程简介、目录、资料、实验指导。")
    add_numbered(doc, "点击章节开始学习，可播放视频、阅读文档或下载附件。")
    add_numbered(doc, "学习进度自动保存，下次登录可从上次离开处继续。")

    add_heading(doc, "4.5  AI 辅导", level=2)
    add_paragraph(doc, "AI 辅导模块提供基于大模型的智能问答与错题解析。")
    add_numbered(doc, "点击“AI 辅导”进入对话页面。")
    add_numbered(doc, "在输入框中输入华为 ICT 相关技术问题，如“OSPF 邻接关系建立过程”。")
    add_numbered(doc, "点击“发送”或按回车键，AI 将返回结构化回答。")
    add_numbered(doc, "在题库练习中答错题目时，系统会自动调用 AI 解析，提供知识点定位与记忆技巧。")
    add_numbered(doc, "若本地 Ollama 模型不可用，系统将自动降级为本地结构化模板解析。")

    add_heading(doc, "4.6  云端实验", level=2)
    add_paragraph(doc, "云端实验模块提供在线实验环境，支持华为 ICT 相关技术实操。")
    add_numbered(doc, "点击“云实验”查看实验列表。")
    add_numbered(doc, "选择实验并查看实验目标、环境说明、操作步骤。")
    add_numbered(doc, "点击“进入实验环境”开始实验。")
    add_numbered(doc, "按照实验指导书完成操作，可在线提交实验报告。")
    add_numbered(doc, "实验进度与成绩将同步至教师端与管理员端。")

    add_heading(doc, "4.7  个人中心与认证", level=2)
    add_paragraph(doc, "个人中心用于管理个人信息与学习认证进度。")
    add_numbered(doc, "点击右上角头像或“个人中心”。")
    add_numbered(doc, "查看并编辑个人资料（姓名、学号、联系方式等）。")
    add_numbered(doc, "查看已获认证、认证进度与推荐认证路径。")
    add_numbered(doc, "可查看历史练习、考试、实验记录及学习统计图表。")
    doc.add_page_break()

    # 5 教师端
    add_heading(doc, "5  教师端操作手册", level=1)
    add_paragraph(doc, "教师端面向任课教师，提供课程管理、学员管理、考试记录、学情分析等功能。")

    add_heading(doc, "5.1  教学仪表盘", level=2)
    add_paragraph(doc, "登录后进入教师端仪表盘，展示所授课程的整体情况。")
    add_bullet(doc, "课程概览：显示课程数量、学员总数、平均完成率。")
    add_bullet(doc, "考试概况：显示学员模拟考试次数、合格率、平均分。")
    add_bullet(doc, "待办事项：显示学员作业、实验报告、注册审核等提醒。")
    add_bullet(doc, "快捷入口：可进入“我的课程”“学员管理”“考试记录”“学情分析”。")

    add_heading(doc, "5.2  我的课程", level=2)
    add_paragraph(doc, "教师可在“我的课程”中管理所授课程。")
    add_numbered(doc, "点击左侧导航栏“我的课程”。")
    add_numbered(doc, "课程列表展示课程名称、认证方向、学员人数、更新状态。")
    add_numbered(doc, "点击“新增课程”填写课程信息、上传封面、添加课程目录与资料。")
    add_numbered(doc, "点击已有课程可编辑课程详情、章节、实验指导。")
    add_numbered(doc, "可设置课程可见范围与学习权限。")

    add_heading(doc, "5.3  学员管理", level=2)
    add_paragraph(doc, "学员管理用于查看与管理所选课程的学生名单。")
    add_numbered(doc, "点击“学员管理”。")
    add_numbered(doc, "学员列表展示姓名、学号、所属课程、学习进度、最近登录时间。")
    add_numbered(doc, "支持按姓名、学号、课程筛选学员。")
    add_numbered(doc, "点击学员姓名可查看其学习详情、练习记录、考试记录、实验成绩。")
    add_numbered(doc, "可对学员进行分组、标记学习状态或发送通知。")

    add_heading(doc, "5.4  考试记录", level=2)
    add_paragraph(doc, "考试记录与学生端模拟考试数据实时同步，教师可查看学员考试情况。")
    add_numbered(doc, "点击“考试记录”。")
    add_numbered(doc, "页面顶部显示考试记录总数、已完成数、进行中数、合格率。")
    add_numbered(doc, "使用搜索框、状态筛选、考试名称筛选定位目标记录。")
    add_numbered(doc, "记录卡片展示学员姓名、考试名称、正确题数、考试时间、完成状态、得分百分比及合格状态。")
    add_numbered(doc, "点击“查看详情”弹出考试详情弹窗，展示学员得分、正确题数、状态、考试时长、开始与完成时间。")
    add_numbered(doc, "在详情弹窗中可查看每道题的答题明细，包括题目内容、选项、正确答案、学生答案、作答状态。")
    add_numbered(doc, "点击“刷新”按钮可重新加载最新考试数据。")

    add_heading(doc, "5.5  学情分析", level=2)
    add_paragraph(doc, "学情分析提供数据可视化，帮助教师掌握班级整体学习情况。")
    add_numbered(doc, "点击“学情分析”。")
    add_numbered(doc, "查看整体数据：学习活跃度、课程完成率、考试合格率、知识点掌握分布。")
    add_numbered(doc, "选择具体课程或班级，查看其学习趋势与薄弱环节。")
    add_numbered(doc, "根据分析结果调整教学计划与重点讲解内容。")
    doc.add_page_break()

    # 6 管理员端
    add_heading(doc, "6  管理员端操作手册", level=1)
    add_paragraph(doc, "管理员端是系统运维核心，负责平台用户、资源、数据、配置的全局管理。")

    add_heading(doc, "6.1  管理仪表盘", level=2)
    add_paragraph(doc, "管理员登录后进入管理仪表盘，展示平台运行全貌。")
    add_bullet(doc, "核心统计：注册学生数、在线课程数、认证通过数、云端实训实例数。")
    add_bullet(doc, "今日概况：今日在线人数、新增注册、刷题次数、认证通过人数。")
    add_bullet(doc, "最近登录用户：展示最近登录的用户名称、角色、登录时间、在线状态。")
    add_bullet(doc, "系统通知与待办：展示新用户审核、课程更新、服务器状态等提醒。")

    add_heading(doc, "6.2  用户管理", level=2)
    add_paragraph(doc, "用户管理用于维护系统所有账户。")
    add_numbered(doc, "点击“用户管理”。")
    add_numbered(doc, "用户列表展示账号、姓名、角色、状态、注册时间。")
    add_numbered(doc, "支持新增用户、批量导入、编辑资料、重置密码、启用/禁用账号。")
    add_numbered(doc, "处理新用户注册审核，通过或拒绝注册申请。")
    add_numbered(doc, "可按角色（学生、教师、管理员）筛选用户。")

    add_heading(doc, "6.3  课程管理", level=2)
    add_paragraph(doc, "课程管理用于维护平台所有课程资源。")
    add_numbered(doc, "点击“课程管理”。")
    add_numbered(doc, "查看课程列表，支持按认证方向、状态、任课教师筛选。")
    add_numbered(doc, "新增课程：填写课程名称、认证方向、课程简介、封面、章节、资料。")
    add_numbered(doc, "编辑课程：修改课程信息、更新章节内容、调整可见状态。")
    add_numbered(doc, "删除课程：对于不再使用的课程，可进行下架或删除操作。")

    add_heading(doc, "6.4  认证管理", level=2)
    add_paragraph(doc, "认证管理用于维护华为 ICT 认证方向与证书信息。")
    add_numbered(doc, "点击“认证管理”。")
    add_numbered(doc, "维护认证方向列表（如 HCIA-Datacom、HCIP-Security、HCIE-Cloud 等）。")
    add_numbered(doc, "设置认证对应的课程、题库、实验与推荐学习路径。")
    add_numbered(doc, "查看学员认证通过情况与统计报表。")

    add_heading(doc, "6.5  考试题库", level=2)
    add_paragraph(doc, "考试题库用于管理模拟考试与题库练习的题目。")
    add_numbered(doc, "点击“考试题库”。")
    add_numbered(doc, "题库列表展示题目、题型、知识点、难度、所属认证方向。")
    add_numbered(doc, "支持单题新增、批量导入（Excel/Word 模板）、编辑、删除。")
    add_numbered(doc, "设置题目属性：题目类型（单选、多选、判断）、选项、正确答案、解析、知识点标签。")
    add_numbered(doc, "组卷管理：将题目组合成模拟考试，设置考试时长、及格分数、题目数量。")
    doc.add_page_break()

    add_heading(doc, "6.6  实训管理", level=2)
    add_paragraph(doc, "实训管理用于维护云端实验环境与实验内容。")
    add_numbered(doc, "点击“实训管理”。")
    add_numbered(doc, "管理实验环境：配置实验镜像、资源配额、实验时长。")
    add_numbered(doc, "维护实验列表：新增、编辑、删除实验项目。")
    add_numbered(doc, "查看学生实验报告与评分。")
    add_numbered(doc, "监控云端实训实例运行状态。")

    add_heading(doc, "6.7  赛事管理", level=2)
    add_paragraph(doc, "赛事管理用于维护华为 ICT 大赛或相关赛事信息。")
    add_numbered(doc, "点击“赛事管理”。")
    add_numbered(doc, "新增赛事：填写赛事名称、时间、报名规则、比赛说明。")
    add_numbered(doc, "管理报名名单与参赛队伍。")
    add_numbered(doc, "发布赛事通知与成绩。")

    add_heading(doc, "6.8  就业管理", level=2)
    add_paragraph(doc, "就业管理用于发布与维护就业相关信息。")
    add_numbered(doc, "点击“就业管理”。")
    add_numbered(doc, "新增就业岗位：填写岗位名称、企业、岗位要求、联系方式。")
    add_numbered(doc, "管理岗位上下架状态。")
    add_numbered(doc, "查看学生投递或关注情况。")

    add_heading(doc, "6.9  数据统计", level=2)
    add_paragraph(doc, "数据统计提供平台级数据报表。")
    add_numbered(doc, "点击“数据统计”。")
    add_numbered(doc, "查看用户增长、活跃度、课程学习、考试练习、实验完成等多维度报表。")
    add_numbered(doc, "支持按时间范围（日、周、月、年）筛选数据。")
    add_numbered(doc, "可导出统计图表与数据表格，用于汇报与分析。")

    add_heading(doc, "6.10  系统设置", level=2)
    add_paragraph(doc, "系统设置用于配置平台基础参数与运行策略。")
    add_numbered(doc, "点击“系统设置”。")
    add_numbered(doc, "配置站点名称、LOGO、版权信息、联系方式。")
    add_numbered(doc, "设置 AI 辅导参数（如 Ollama 模型地址、备用解析开关）。")
    add_numbered(doc, "配置注册审核策略、登录安全策略、密码规则。")
    add_numbered(doc, "管理系统公告与消息模板。")
    doc.add_page_break()

    # 7 常见问题
    add_heading(doc, "7  常见问题与注意事项", level=1)
    add_heading(doc, "7.1  登录问题", level=2)
    add_bullet(doc, "若提示“用户名或密码错误”，请检查大小写与角色入口是否正确。")
    add_bullet(doc, "若账号被禁用，请联系管理员确认状态。")
    add_bullet(doc, "建议使用 Chrome、Edge、Firefox 等主流浏览器访问。")

    add_heading(doc, "7.2  考试与练习问题", level=2)
    add_bullet(doc, "模拟考试开始后倒计时无法暂停，请确保网络稳定、时间充足。")
    add_bullet(doc, "考试过程中如遇页面异常刷新，重新登录后可在“考试记录”中查看是否已保存进度。")
    add_bullet(doc, "题库练习中的 AI 解析需依赖本地 Ollama 服务，若服务不可用，将自动降级为模板解析。")

    add_heading(doc, "7.3  数据同步说明", level=2)
    add_bullet(doc, "学生端模拟考试记录会实时同步到教师端“考试记录”。")
    add_bullet(doc, "课程学习进度、实验报告、练习数据均会实时更新至相关统计页面。")

    add_heading(doc, "7.4  安全与隐私", level=2)
    add_bullet(doc, "请勿将个人账号密码告知他人。")
    add_bullet(doc, "离开工位时请主动退出系统。")
    add_bullet(doc, "管理员应定期修改后台管理密码，并妥善保管敏感操作权限。")

    add_heading(doc, "7.5  联系与支持", level=2)
    add_paragraph(doc, "如在使用过程中遇到无法解决的问题，请联系系统管理员或技术支持人员，并提供以下信息：")
    add_bullet(doc, "问题发生的页面与操作步骤。")
    add_bullet(doc, "浏览器版本与操作系统信息。")
    add_bullet(doc, "错误提示截图或文字描述。")

    # 保存
    doc.save(str(OUTPUT_PATH))
    print(f"操作手册已生成：{OUTPUT_PATH}")

if __name__ == '__main__':
    create_manual()
